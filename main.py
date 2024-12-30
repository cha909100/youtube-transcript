import os
import csv
import re
import time
from datetime import datetime
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document
import openai

# Set up OpenAI API key
openai.api_key = os.environ.get("OPENAI_API_KEY")

def get_channel_id(youtube, channel_url):
    """Extract the channel ID from a YouTube URL."""
    if "channel/" in channel_url:
        return channel_url.split("channel/")[-1]
    elif "user/" in channel_url:
        username = channel_url.split("user/")[-1]
        try:
            response = youtube.channels().list(forUsername=username, part="id").execute()
            return response["items"][0]["id"] if response["items"] else None
        except HttpError as e:
            print(f"Error fetching channel ID: {e}")
            return None
    return None

def get_channel_details(youtube, channel_id):
    """Retrieve details of a YouTube channel by ID."""
    try:
        response = youtube.channels().list(id=channel_id, part="snippet,contentDetails").execute()
        if not response.get('items'):
            return None
        return response['items'][0]
    except HttpError as e:
        print(f"Error fetching channel details: {e}")
        return None

def list_all_videos(api_key, channel_id):
    """List all videos from a YouTube channel and retrieve their transcripts and statistics."""
    youtube = build('youtube', 'v3', developerKey=api_key)
    videos = []
    count = 1

    channel_info = get_channel_details(youtube, channel_id)
    if not channel_info:
        return None, "Channel not found"

    channel_name = channel_info['snippet']['title']
    uploads_playlist_id = channel_info['contentDetails']['relatedPlaylists']['uploads']

    request = youtube.playlistItems().list(playlistId=uploads_playlist_id, part='snippet,contentDetails', maxResults=50)
    while request:
        response = request.execute()
        video_ids = [item['contentDetails']['videoId'] for item in response['items']]
        video_details_response = youtube.videos().list(id=','.join(video_ids), part='snippet,statistics').execute()
        for item in video_details_response['items']:
            video_id = item['id']
            video_url = f"https://www.youtube.com/watch?v={video_id}"
            video_title = item['snippet']['title']
            video_date = datetime.strptime(item['snippet']['publishedAt'], '%Y-%m-%dT%H:%M:%SZ').strftime('%Y-%m-%d')
            likes = item['statistics'].get('likeCount', 'No likes data')
            views = item['statistics'].get('viewCount', 'No views data')
            try:
                transcript = YouTubeTranscriptApi.get_transcript(video_id)
                combined_text = ' '.join([entry['text'] for entry in transcript])
                summary, points = summarize_and_extract_points(combined_text)
                videos.append((count, video_title, video_url, video_date, likes, views, summary, points))
            except Exception:
                videos.append((count, video_title, video_url, video_date, likes, views, "No transcript available", []))
            count += 1
            time.sleep(0.5)  # Rate limiting to avoid quota issues
        request = youtube.playlistItems().list_next(request, response)
    return videos, channel_name

def summarize_and_extract_points(text):
    """Generate a summary and key discussion points from a text using OpenAI's GPT model."""
    try:
        response = openai.Completion.create(
            prompt=f"Summarize this text and provide 3-5 key discussion points: {text}",
            model="text-davinci-003",
            max_tokens=150
        )
        result = response.choices[0].text.strip()
        lines = result.split('\n')
        summary = lines[0]
        points = [line for line in lines[1:] if line.strip()]
        return summary, points
    except Exception as e:
        print(f"Error summarizing text: {e}")
        return "Summary unavailable", []

def save_videos_to_csv(videos, filename):
    """Save video data to a CSV file."""
    with open(filename, 'w', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)
        writer.writerow(['No.', 'Video Title', 'Video URL', 'Upload Date', 'Likes', 'Views', 'Summary', 'Discussion Points'])
        for video in videos:
            writer.writerow(video)

def save_summary_to_docx(videos, directory="summaries", prefix="OPENAI"):
    """Save video summaries and discussion points to DOCX files with a specified prefix."""
    if not os.path.exists(directory):
        os.makedirs(directory)
    for video in videos:
        doc = Document()
        doc.add_heading(video[1], 0)
        doc.add_paragraph(video[6])
        doc.add_heading('Discussion Points', level=1)
        for point in video[7]:
            doc.add_paragraph(point, style='ListBullet')
        filename = f"{prefix}_{re.sub(r'[\\/*?:"<>|]', '', video[1])}.docx"
        doc.save(os.path.join(directory, filename))

if __name__ == "__main__":
    youtube_api_key = input("Please enter your YouTube API key: ").strip()
    channel_url = input("Please enter the YouTube channel URL: ").strip()

    youtube = build("youtube", "v3", developerKey=youtube_api_key)
    channel_id = get_channel_id(youtube, channel_url)

    if not channel_id:
        print("Invalid channel URL or channel ID not found.")
        exit()

    videos, channel_name = list_all_videos(youtube_api_key, channel_id)
    if videos:
        csv_filename = f"{channel_name.replace(' ', '_')}.csv"
        save_videos_to_csv(videos, csv_filename)
        save_summary_to_docx(videos)
        print(f"Data saved to {csv_filename} and DOCX summaries.")
    else:
        print("No videos found or channel not accessible.")
